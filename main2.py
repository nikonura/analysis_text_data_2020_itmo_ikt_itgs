import docx
import re


def algorithm(cell, block_separators, block, paragraphs, word_matching_list):
	for i in block_separators:
		expression = block_separators[i]
		if re.findall(expression, cell.text):
			if len(block) > 0:
				paragraphs.append(block)
			block = []
	for i in word_matching_list.keys():
		expression = word_matching_list[i]
		node_matches = re.findall(expression, cell.text)
		if len(node_matches) > 0:
			for j in range(len(node_matches)):
				cell.text = cell.text.replace(node_matches[j], '<'+i+'>')
	block.append(cell.text)
	return block, paragraphs


def get_doc_paragraphs(path, save_path):
	document = docx.Document(docx_path)
	paragraphs = []
	block = []

	word_matching_list = {'sex': r'[Пп]ол[: ]? \w+\b',
						'age': r'[Вв]озраст[: ]? \d{1,3}лет',
						'years_old': r'\d{1,3}[ ]?год[а]?[ ]?\d{0,2}[ ]?[месяц]?[ев]?[ ]?\d{0,2}',
						'birth_place': r'[Мм]есто рождения[: ]? .*',
						'person': r'личность[: ]? .*',
						'registration': r'[Рр]егистраци.[: ]? .*',
						'insurance': r'страхов.+[: ]? .*',
						'document': r'(СНИЛС|ИНН|[Пп]аспорт)[: ]? .*',
						'nationality': r'[Гг]ражданство[: ]? .*',
						'family_status': r'[Сс]емейное положение[: ]? (замужем|женат)',
						'income': r'[Дд]оход[: ]? .*',
						'living_place': r'[Мм]есто жительства[: ]? .*',
						'phone': r'[Тт]елефон[у]?[:\s]+\d{3,}\b',
						'email': r'[A-z\w]@[A-z].[a-z]',
						'floor1': r'(\d? Этаж \d?)',
						'floor2': r'(\d? этаж \d?.*$)',
						'floor3': r'(\d{1,2}эт)',
						'reg_num': r'/?[0-9]{5,}/?',
						'alphanum_code': r'[A-ZА-Я]\d{2,}',
						# 'alphanum_code2': r'[A-ZА-Я]?\d+[A-ZА-Я]+\d+[A-ZА-Я]+)',
						# 'period': r'\d{2}[.]\d{2}[.]\d{2,4}[ -]+\d{2}[.]\d{2}.\d{2}',
						# 'date_time': r'\d{2}[.]\d{2}[.]\d{2,4}[ ]?\d{1,2}[:]\d{2}[:]?\d{2}?',
						'date': r'\d{2}[.]\d{2}[.]\d{2,4}',
						'time_full': r'\d{1,2}[:]\d{2}[:]\d{2}',
						'time': r'\d{1,2}[:]\d{2}',
						'blood_pressure': r'\d{2,3}[\\]?[/]?\d{2,3} мм [рp]т[.]?[ ]?[сc]т[.]?',
						'nums_with_sep': r'\d{1,3}\.\d{1,3}',
						'FIO_short_double': r'[А-ЯЁ][а-яё]+[ ]+[А-ЯЁ]{2}[ \t\n\r.?!/$]+',
						'FIO_short_with_space': r'[А-ЯЁ][а-яё]+\s+[А-ЯЁ]{1}[.]?[ ]+[А-ЯЁ]{1}[.]?[ \t\n\r.?!$]+',
						'FIO_short_without_space': r'[А-ЯЁ][а-яё]+\s+[А-ЯЁ]{1}[.]?[А-ЯЁ]{1}[.]?[ \t\n\r.?!$]+',
						'FIO_full': r'[А-ЯЁ][а-яё\-]+[ ]+[А-ЯЁ][а-яё\-]+[ ]*[А-ЯЁ][а-яё\-]+',
						'abbreviation': r'[ ,.][А-ЯЁA-Z]{3,4}\b',
						'place': r'>.+ Минздрава России'
						}

	block_separators = {
						'2slash1': r'[\w -]+ /+ \d{2}[.]\d{2}[.]\d{2,4} \d{0,2}[:.]?\d{0,2}[ ]?/+ [\w -]+',
						'2slash2': r'\d{2}[.]\d{2}[.]\d{2,4} \d{0,2}[:.]?\d{0,2}[ ][/][ \w]+[/][ \w]+',
						'1slash': r'\d{2}[.]\d{2}[.]\d{2,4} \d{0,2}[:.]?\d{0,2}[ ]?/+',
						'capital_letters4': r'[А-ЯЁ]+\b[ ]?[А-ЯЁ]+\b[ ]+[А-ЯЁ]+\b[ ]+[А-ЯЁ]+\b',
						'capital_letters3': r'[А-ЯЁ]{5,}\b[ ]?[А-ЯЁ]+\b[ ]?[А-ЯЁ]{5,}\b',
						'capital_letters2': r'[А-ЯЁ]{5,}\b[ ]?[А-ЯЁ]{5,}\b',
						'capital_letters1': r'^[А-ЯЁ]{5,}\b$',
						}

	# Update header
	for sec in document.sections:
		if len(sec.header.paragraphs)==0:
			for table in sec.header.tables:
				for row in table.rows:
					for cell in row.cells:
						block, paragraphs = algorithm(cell, block_separators, block, paragraphs, word_matching_list)
		else:
			for para in sec.header.paragraphs:
				block, paragraphs = algorithm(para, block_separators, block, paragraphs, word_matching_list)
	if len(block) > 0:
		paragraphs.append(block)

	# Update footer
	for sec in document.sections:
		if len(sec.footer.paragraphs)==0:
			for table in sec.footer.tables:
				for row in table.rows:
					for cell in row.cells:
						block, paragraphs = algorithm(cell, block_separators, block, paragraphs, word_matching_list)
		else:
			for para in sec.footer.paragraphs:
				block, paragraphs = algorithm(para, block_separators, block, paragraphs, word_matching_list)
	if len(block) > 0:
		paragraphs.append(block)

	# Update body
	texts = document._element.xpath('//w:t')
	for n in range(len(texts)):
		block, paragraphs = algorithm(texts[n], block_separators, block, paragraphs, word_matching_list)
	if len(block) > 0:
		paragraphs.append(block)

	document.save(save_path)
	return paragraphs


# docx_path = "/home/nikon-cook/Documents/МИТМО/Analisys_TD/Med_karta_1_bez_personalnykh_dannykh.docx"
# save_path = "/home/nikon-cook/Documents/МИТМО/Analisys_TD/Med_karta_1_bez_personalnykh_dannykh_new.docx"
docx_path = "Med_karta_1_bez_personalnykh_dannykh.docx"
save_path = "Med_karta_1_bez_personalnykh_dannykh_new.docx"
# docx_path = "medical_record.docx"
# save_path = "medical_record_new.docx"

paragraphs = get_doc_paragraphs(docx_path, save_path)
print(len(paragraphs))

# for i in paragraphs:
# 	print(i)
