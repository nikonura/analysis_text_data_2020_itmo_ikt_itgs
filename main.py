import docx
import re

changes = {}

word_matching_list = {
	'sex': r'\b[Пп]ол[: ]? \w+\b',
	'years_old': r'\d{1,3}[ ]?[гл][ое][дт][а]?|\d{1,2}[ ]?месяц[е]?[в]?',
	'age': r'[Вв]озраст[: ]? \d{1,3}',
	'birth_place': r'[Мм]есто рождения[: ]? .*',
	'person': r'личность[: ]? .*',
	'registration': r'[Рр]егистраци.[: ]? .*',
	'insurance': r'страхов.+[: ]? .*',
	'document': r'(СНИЛС[: ]? .*|ИНН[: ]? .*|[Пп]аспорт[: ]? .*)',
	'nationality': r'[Гг]ражданство[: ]? .*',
	'family_status': r'[Сс]емейное положение[: ]? (замужем|женат)',
	'income': r'[Дд]оход[: ]? .*',
	'living_place': r'[Мм]есто жительства[: ]? .*|[Аа]дрес.[: ]? .*',
	'phone': r'[Тт]елефон[у]?[:\s]+\d{3,}\b',
	'email': r'[A-z\w]@[A-z].[a-z]',
	'floor1': r'(\d? Этаж \d?)',
	'floor2': r'(\d? этаж \d?.*$)',
	'floor3': r'(\d{1,2}эт)',
	'reg_num': r'/?[0-9]{5,}/?',
	'alphanum_code': r'[A-ZА-Я]\d{2,}',
	# 'alphanum_code2': r'[A-ZА-Я]?\d+[A-ZА-Я]+\d+[A-ZА-Я]+)',
	# 'period': r'\d{2}[.]\d{2}[.]\d{2,4}[ -]+\d{2}[.]\d{2}.\d{2}',
	# 'date_time': r'\d{2}[.]\d{2}[.]\d{2,4}[ ]?\d{1,2}[:]\d{2}[:]?\d{2}?',
	'date': r'\d{1,2}[.]\d{2}[.]\d{2,4}',
	'time_full': r'\d{1,2}[:]\d{2}[:]\d{2}',
	'time': r'\d{1,2}[:]\d{2}',
	'blood_pressure': r'\d{2,3}[\\]?[/]?\d{2,3} мм [рp]т[.]?[ ]?[сc]т[.]?',
	# 'nums_with_sep': r'\d{1,3}\.\d{1,3}',
	'FIO_short_double': r'[А-ЯЁ][а-яё]+[ ]+[А-ЯЁ]{2}',
	'FIO_short_with_space': r'[А-ЯЁ][а-яё]+\s+[А-ЯЁ]{1}[.]?[ ]+[А-ЯЁ]{1}[.]?[ \t\n\r.?!$]+',
	'FIO_short_without_space': r'[А-ЯЁ][а-яё]+\s+[А-ЯЁ]{1}[.]?[А-ЯЁ]{1}[.]?[ \t\n\r.?!$]+',
	'FIO_full': r'[А-ЯЁ][а-яё\-]+[ ]+[А-ЯЁ][а-яё\-]+[ ]*[А-ЯЁ][а-яё\-]+',
	# 'abbreviation': r'[ ,.][А-ЯЁA-Z]{3,4}\b',
	'place': r'>.+ Минздрава России'
}

block_separators = {
	'2slash1': r'[\w -]+ /+ \d{2}[.]\d{2}[.]\d{2,4} \d{0,2}[:.]?\d{0,2}[ ]?/+ [\w -]+',
	'2slash2': r'\d{2}[.]\d{2}[.]\d{2,4} \d{0,2}[:.]?\d{0,2}[ ][/][ \w]+[/][ \w]+',
	'1slash': r'\d{2}[.]\d{2}[.]\d{2,4} \d{0,2}[:.]?\d{0,2}[ ]?/+',
	'capital_letters4': r'[А-ЯЁ]+\b[ ]?[А-ЯЁ]+\b[ ]+[А-ЯЁ]+\b[ ]+[А-ЯЁ]+\b',
	'capital_letters3': r'[А-ЯЁ]{5,}\b[ ]?[А-ЯЁ]+\b[ ]?[А-ЯЁ]{5,}\b',
	'capital_letters2': r'[А-ЯЁ]{5,}\b[ ]?[А-ЯЁ]{5,}\b',
	'capital_letters1': r'^[А-ЯЁ]{5,}\b$',
}


def algorithm(cell, block, paragraphs):
	global word_matching_list
	global block_separators
	global changes
	# print('\n', cell.text)  # для демонстрации
	for i in block_separators:
		expression = block_separators[i]
		if re.findall(expression, cell.text):
			if len(block) > 0:
				paragraphs.append(block)
			block = []
	for i in word_matching_list.keys():
		expression = word_matching_list[i]
		node_matches = re.findall(expression, cell.text)
		if len(node_matches) > 0:
			# print(i, node_matches)  # для демонстрации
			for j in range(len(node_matches)):
				dict_len = len(changes.keys()) + 1
				changes['< ' + i + ' ' + str(dict_len) + ' >'] = node_matches[j]
				cell.text = cell.text.replace(node_matches[j], '<' + i + '>')
	block.append(cell.text)
	return block, paragraphs


def get_doc_paragraphs(path, save_path):
	global word_matching_list
	global block_separators
	document = docx.Document(docx_path)
	paragraphs = []
	block = []

	# Update header
	for sec in document.sections:
		if len(sec.header.paragraphs) == 0:
			for table in sec.header.tables:
				for row in table.rows:
					for cell in row.cells:
						block, paragraphs = algorithm(cell, block, paragraphs)
		else:
			for para in sec.header.paragraphs:
				block, paragraphs = algorithm(para, block, paragraphs)
	if len(block) > 0:
		paragraphs.append(block)

	# Update footer
	for sec in document.sections:
		if len(sec.footer.paragraphs) == 0:
			for table in sec.footer.tables:
				for row in table.rows:
					for cell in row.cells:
						block, paragraphs = algorithm(cell, block, paragraphs)
		else:
			for para in sec.footer.paragraphs:
				block, paragraphs = algorithm(para, block, paragraphs)
	if len(block) > 0:
		paragraphs.append(block)

	# Update body
	if len(document.paragraphs) != 0:
		for para in document.paragraphs:
			block, paragraphs = algorithm(para, block, paragraphs)
	texts = document._element.xpath('//w:t')
	for n in range(len(texts)):
		block, paragraphs = algorithm(texts[n], block, paragraphs)
	if len(block) > 0:
		paragraphs.append(block)

	document.save(save_path)
	return paragraphs


# filename options
# option = "Med_karta_1_bez_personalnykh_dannykh"
option = "Med_karta_2"

# docx_path = "/home/nikon-cook/Documents/МИТМО/Analisys_TD/Med_karta_1_bez_personalnykh_dannykh.docx"
# save_path = "/home/nikon-cook/Documents/МИТМО/Analisys_TD/Med_karta_1_bez_personalnykh_dannykh_new.docx"
docx_path = option + ".docx"
save_path = option + "_new.docx"


paragraphs = get_doc_paragraphs(docx_path, save_path)

def print_all(changes_if, para_if):
	global changes
	global paragraphs

	print('Paragraphs count:', len(paragraphs))
	if changes_if == 1:
		for i in changes.keys():
			print(i, changes[i])
		print('\n')

	print('Changes dictionary:', len(changes.keys()))
	if para_if == 1:
		for i in paragraphs:
			print(i)
		print('\n')

print_all(0, 0)
print()


# F-мера для карты 2

fp = 1  # Соли Ж.К.
fn = 2  # Даты в формате d.mm или dd.mm
tp = len(changes.keys()) - fp  # 35

recall = tp / (tp + fn)
precision = tp / (tp + fp)
f_score = 2 * precision * recall / (precision + recall)

print(f'TP={tp}, FP={fp}, FN={fn}')
print('Recall:', recall)
print('Precision:', precision)
print('F-мера:', f_score)
